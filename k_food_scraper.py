import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
import re
import time

SOURCES = [
    "https://en.wikipedia.org/wiki/Korean_cuisine",
    "https://www.korea.net/AboutKorea/Culture-and-the-Arts/K-food",
    "https://english.visitkorea.or.kr/enu/FO/FO_EN_15.jsp"
]

HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}

def clean_text(s):
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def first_paragraph(soup, min_len=120):
    for p in soup.find_all('p'):
        text = clean_text(p.get_text())
        if len(text) >= min_len:
            return text
    # fallback: first non-empty paragraph
    for p in soup.find_all('p'):
        text = clean_text(p.get_text())
        if text:
            return text
    return ""

def fetch_summary(url):
    try:
        r = requests.get(url, headers=HEADERS, timeout=10)
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        title = (soup.title.string or "").strip() if soup.title else url
        summary = first_paragraph(soup)
        if not summary:
            # try meta description
            md = soup.find("meta", attrs={"name":"description"}) or soup.find("meta", attrs={"property":"og:description"})
            if md and md.get("content"):
                summary = clean_text(md.get("content"))
        # truncate safely
        if len(summary) > 1000:
            summary = summary[:1000].rsplit(' ', 1)[0] + "..."
        return {"title": title, "summary": summary, "url": url}
    except Exception as e:
        return {"title": url, "summary": f"Failed to fetch: {e}", "url": url}

def make_doc(items, out_path=r"c:\work\k-food.docx"):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading("K-food 자료 모음", level=0)
    for it in items:
        doc.add_heading(it['title'], level=1)
        doc.add_paragraph(it['summary'])
        doc.add_paragraph("Source: " + it['url'])
        doc.add_paragraph()  # spacer

    doc.save(out_path)
    print("Saved:", out_path)

def main():
    results = []
    for url in SOURCES:
        print("Fetching:", url)
        item = fetch_summary(url)
        results.append(item)
        time.sleep(1)  # polite delay
    make_doc(results)

if __name__ == "__main__":
    main()